import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class ParsedDocument:
    title: str = ""
    abstract: str = ""
    body: str = ""
    references: str = ""
    full_text: str = ""
    source_path: str = ""


_REF_HEADERS = re.compile(
    r"^\s*(references|bibliography|works cited|reference list)\s*$",
    re.IGNORECASE | re.MULTILINE,
)


def _split_references(text: str) -> tuple[str, str]:
    match = _REF_HEADERS.search(text)
    if match:
        return text[: match.start()].strip(), text[match.start() :].strip()
    return text.strip(), ""


def _extract_abstract(text: str) -> tuple[str, str]:
    m = re.search(r"\bAbstract\b[\s\S]{0,20}?\n([\s\S]{50,1500}?)(\n\s*\n|\b(Introduction|1\.)\b)", text, re.IGNORECASE)
    if m:
        return m.group(1).strip(), text
    return "", text


def parse_pdf(path: str) -> ParsedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    pages = [page.get_text() for page in doc]
    doc.close()

    full_text = "\n".join(pages)
    body_and_refs, references = _split_references(full_text)
    abstract, _ = _extract_abstract(body_and_refs)

    title = ""
    for line in full_text.splitlines():
        line = line.strip()
        if 10 < len(line) < 200 and not line.lower().startswith("abstract"):
            title = line
            break

    return ParsedDocument(
        title=title,
        abstract=abstract,
        body=body_and_refs,
        references=references,
        full_text=full_text,
        source_path=path,
    )


def parse_docx(path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)

    title = ""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") or p.style.name == "Title":
            title = p.text.strip()
            break
    if not title:
        for line in full_text.splitlines():
            line = line.strip()
            if 10 < len(line) < 200:
                title = line
                break

    body_and_refs, references = _split_references(full_text)
    abstract, _ = _extract_abstract(body_and_refs)

    return ParsedDocument(
        title=title,
        abstract=abstract,
        body=body_and_refs,
        references=references,
        full_text=full_text,
        source_path=path,
    )


def parse(path: str) -> ParsedDocument:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
